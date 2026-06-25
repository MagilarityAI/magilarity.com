"""
report_formatter.py — Форматування результатів аналізу ТД у DOCX файли.

Отримує результати base_analyzer та contract_analyzer (analysis dict + contract_analysis dict)
і генерує 4 вихідних DOCX файли + зберігає analysis.json.

Файли виводу:
  1. analysis_report.docx    — 16+2 розділів, кольорове кодування
  2. documents_checklist.docx — чеклист документів пропозиції (6 колонок)
  3. contract_analysis.docx  — юридичний аналіз проекту договору
  4. winner_checklist.docx   — чеклист переможця (3 розділи)

Модель: claude-sonnet-4-6 (для структурування/збагачення де потрібно)
"""

import json
import logging
import re
import time
from datetime import datetime, date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Кольори ────────────────────────────────────────────────────────────────────

COLOR_RED = RGBColor(0xC0, 0x00, 0x00)       # критично
COLOR_ORANGE = RGBColor(0xE2, 0x6B, 0x0A)    # ризик
COLOR_GREEN = RGBColor(0x37, 0x86, 0x36)     # норма
COLOR_DARK = RGBColor(0x1F, 0x1F, 0x1F)      # заголовки
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)      # примітки
COLOR_BLUE = RGBColor(0x1F, 0x49, 0x7D)      # посилання/вердикт


# ── Допоміжні функції стилю ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Встановлює колір фону комірки таблиці."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Додає заголовок відповідного рівня."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = COLOR_DARK


def _add_paragraph(doc: Document, text: str, bold: bool = False,
                   color: Optional[RGBColor] = None, size_pt: int = 11) -> None:
    """Додає параграф з опціональним форматуванням."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _verdict_color(verdict: Optional[str]) -> RGBColor:
    """Повертає колір для вердикту."""
    mapping = {
        'recommended': COLOR_GREEN,
        'risky': COLOR_ORANGE,
        'not_recommended': COLOR_RED,
        'not_recommended_without_appeal': COLOR_RED,
    }
    return mapping.get(verdict or '', COLOR_DARK)


def _verdict_ua(verdict: Optional[str]) -> str:
    """Локалізує вердикт для відображення."""
    mapping = {
        'recommended': 'РЕКОМЕНДОВАНО',
        'risky': 'РИЗИКОВАНО (аналізуйте додатково)',
        'not_recommended': 'НЕ РЕКОМЕНДУЄТЬСЯ',
        'not_recommended_without_appeal': 'НЕ РЕКОМЕНДУЄТЬСЯ (без оскарження)',
    }
    return mapping.get(verdict or '', verdict or '—')


def _risk_color(risk_level: Optional[str]) -> RGBColor:
    """Повертає колір для рівня ризику."""
    mapping = {
        'low': COLOR_GREEN,
        'medium': COLOR_ORANGE,
        'high': COLOR_RED,
        'critical': COLOR_RED,
    }
    return mapping.get(risk_level or '', COLOR_DARK)


def _risk_ua(risk_level: Optional[str]) -> str:
    """Локалізує рівень ризику."""
    mapping = {
        'low': 'Низький',
        'medium': 'Середній',
        'high': 'Високий',
        'critical': 'Критичний',
    }
    return mapping.get(risk_level or '', risk_level or '—')


def _balance_ua(balance: Optional[str]) -> str:
    """Локалізує шкалу балансу договору."""
    mapping = {
        'balanced': 'Збалансований',
        'slightly_skewed_to_customer': 'Незначний перекіс на користь замовника',
        'skewed_to_customer': 'Помітний перекіс на користь замовника',
        'heavily_skewed_to_customer': 'Значний перекіс — явно невигідний для переможця',
    }
    return mapping.get(balance or '', balance or '—')


def _format_value(value) -> str:
    """Форматує числове значення як грошову суму."""
    if isinstance(value, (int, float)):
        return f"{value:,.2f}".replace(',', ' ')
    return str(value or '—')


# ── Стандартне форматування документа ─────────────────────────────────────────

def _setup_document() -> Document:
    """Створює Document з базовими налаштуваннями."""
    doc = Document()

    # Поля сторінки
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Базовий шрифт Normal стилю
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    return doc


# ── Шапка документа ───────────────────────────────────────────────────────────

def _add_doc_header(doc: Document, analysis: dict, subtitle: str) -> None:
    """Додає шапку документа із загальною інформацією про закупівлю."""
    # Назва
    title_p = doc.add_heading('АНАЛІЗ ТЕНДЕРНОЇ ДОКУМЕНТАЦІЇ', level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Підзаголовок
    sub_p = doc.add_paragraph(subtitle)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub_p.runs:
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_BLUE

    doc.add_paragraph()

    # Таблиця основних даних
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ('Публічний ID закупівлі:', analysis.get('public_id', '—')),
        ('Замовник:', analysis.get('customer_name', '—')),
        ('ЄДРПОУ:', analysis.get('customer_edrpou', '—')),
        ('Код ДК:', analysis.get('dk_code', '—')),
        ('Очікувана вартість:', f"{_format_value(analysis.get('expected_value'))} грн"),
        ('Дедлайн подачі:', analysis.get('submission_deadline', '—')),
        ('Дата аналізу:', analysis.get('analyzed_at', datetime.now().strftime('%d.%m.%Y %H:%M'))),
    ]

    for i, (label, value) in enumerate(rows_data):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(cells[0], 'D9E1F2')

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. analysis_report.docx
# ═══════════════════════════════════════════════════════════════════════════════

def _build_analysis_report(analysis: dict, contract_analysis: Optional[dict]) -> Document:
    """
    Будує analysis_report.docx — головний аналітичний звіт.

    Структура: шапка → вердикт → порушення → приховані вимоги →
               підстави оскарження → деталі по блоках → договір → Q&A →
               профіль замовника → метадані.
    """
    doc = _setup_document()
    _add_doc_header(doc, analysis, 'АНАЛІТИЧНИЙ ЗВІТ')

    # ── Розділ 1: Загальний вердикт ───────────────────────────────────────────
    _add_heading(doc, '1. ЗАГАЛЬНИЙ ВЕРДИКТ', level=2)

    verdict = analysis.get('verdict')
    risk_level = analysis.get('risk_level')

    verdict_table = doc.add_table(rows=4, cols=2)
    verdict_table.style = 'Table Grid'

    v_rows = [
        ('ВЕРДИКТ:', _verdict_ua(verdict)),
        ('Рівень ризику:', _risk_ua(risk_level)),
        ('Участь рекомендована:', 'Так' if analysis.get('participate') else 'Ні'),
        ('Можливе оскарження:', 'Так' if analysis.get('appeal_possible') else 'Ні'),
    ]

    for i, (label, value) in enumerate(v_rows):
        cells = verdict_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True

        if label == 'ВЕРДИКТ:':
            for run in cells[1].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = _verdict_color(verdict)
        elif label == 'Рівень ризику:':
            for run in cells[1].paragraphs[0].runs:
                run.font.color.rgb = _risk_color(risk_level)

    doc.add_paragraph()

    # Короткий висновок
    short_summary = analysis.get('short_summary')
    if short_summary:
        p = doc.add_paragraph()
        run = p.add_run('Короткий висновок: ')
        run.bold = True
        p.add_run(short_summary)

    # Дедлайн оскарження
    appeal_deadline = analysis.get('appeal_deadline')
    if appeal_deadline:
        p = doc.add_paragraph()
        run = p.add_run(f'⚠️ Дедлайн оскарження умов ТД: {appeal_deadline}')
        run.bold = True
        run.font.color.rgb = COLOR_RED

    doc.add_paragraph()

    # ── Розділ 2: Порушення законодавства ─────────────────────────────────────
    law_violations = analysis.get('law_violations') or []
    _add_heading(doc, f'2. ПОРУШЕННЯ ЗАКОНОДАВСТВА ({len(law_violations)} шт.)', level=2)

    if law_violations:
        viol_table = doc.add_table(rows=1 + len(law_violations), cols=4)
        viol_table.style = 'Table Grid'

        # Заголовок
        headers = ['ID', 'Стаття', 'Опис порушення', 'Можливе оскарження']
        for j, h in enumerate(headers):
            cell = viol_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, '1F497D')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i, v in enumerate(law_violations):
            row = viol_table.rows[i + 1]
            row.cells[0].text = str(v.get('id', '—'))
            row.cells[1].text = str(v.get('law_article', v.get('article', '—')))
            desc = v.get('description', '—')
            explanation = v.get('explanation', '')
            row.cells[2].text = f"{desc}\n{explanation}" if explanation else desc
            appeal = v.get('appeal_possible', False)
            row.cells[3].text = '✓ Так' if appeal else 'Ні'
            if appeal:
                for run in row.cells[3].paragraphs[0].runs:
                    run.font.color.rgb = COLOR_GREEN
    else:
        doc.add_paragraph('Порушень законодавства не виявлено.')

    doc.add_paragraph()

    # ── Розділ 3: Приховані вимоги ────────────────────────────────────────────
    hidden_reqs = analysis.get('hidden_requirements') or []
    _add_heading(doc, f'3. ПРИХОВАНІ ВИМОГИ ({len(hidden_reqs)} шт.)', level=2)

    if hidden_reqs:
        hr_table = doc.add_table(rows=1 + len(hidden_reqs), cols=4)
        hr_table.style = 'Table Grid'

        headers = ['Документ', 'Де захована', 'Цитата з ТД', 'Примітка']
        for j, h in enumerate(headers):
            cell = hr_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, 'E2EFDA')

        for i, req in enumerate(hidden_reqs):
            row = hr_table.rows[i + 1]
            row.cells[0].text = str(req.get('document', '—'))
            row.cells[1].text = str(req.get('hidden_in', req.get('location', '—')))
            row.cells[2].text = str(req.get('td_quote', req.get('quote', '—')))
            row.cells[3].text = str(req.get('risk_note', req.get('note', '')))
    else:
        doc.add_paragraph('Прихованих вимог не виявлено.')

    doc.add_paragraph()

    # ── Розділ 4: Підстави для оскарження ────────────────────────────────────
    appeal_grounds = analysis.get('appeal_grounds') or []
    _add_heading(doc, f'4. ПІДСТАВИ ДЛЯ ОСКАРЖЕННЯ ({len(appeal_grounds)} шт.)', level=2)

    if appeal_grounds:
        if appeal_deadline:
            p = doc.add_paragraph()
            run = p.add_run(f'Дедлайн подачі скарги: {appeal_deadline}')
            run.bold = True
            run.font.color.rgb = COLOR_RED

        for i, g in enumerate(appeal_grounds, 1):
            _add_heading(doc, f'4.{i}. {g.get("ground", "—")}', level=3)

            ground_table = doc.add_table(rows=3, cols=2)
            ground_table.style = 'Table Grid'

            g_rows = [
                ('Правова підстава:', g.get('law_basis', g.get('legal_basis', '—'))),
                ('Цитата з ТД:', g.get('td_quote', '—')),
                ('Запропонована вимога:', g.get('suggested_claim', '—')),
            ]

            for j, (label, value) in enumerate(g_rows):
                cells = ground_table.rows[j].cells
                cells[0].text = label
                cells[1].text = str(value)
                cells[0].paragraphs[0].runs[0].bold = True
                _set_cell_bg(cells[0], 'FCE4D6')

            doc.add_paragraph()
    else:
        doc.add_paragraph('Підстав для оскарження не виявлено.')

    doc.add_paragraph()

    # ── Розділ 5: Гарантійне забезпечення ────────────────────────────────────
    guarantee = analysis.get('guarantee_requirements') or {}
    _add_heading(doc, '5. ВИМОГИ ДО ТЕНДЕРНОГО ЗАБЕЗПЕЧЕННЯ', level=2)

    if guarantee:
        g_table = doc.add_table(rows=len(guarantee), cols=2)
        g_table.style = 'Table Grid'

        label_map = {
            'required': 'Вимагається:',
            'amount': 'Сума:',
            'type': 'Тип:',
            'notes': 'Примітки:',
            'percent': 'Відсоток:',
        }

        for i, (key, value) in enumerate(guarantee.items()):
            cells = g_table.rows[i].cells
            cells[0].text = label_map.get(key, key)
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].bold = True
    else:
        doc.add_paragraph('Вимоги до тендерного забезпечення не зазначені.')

    doc.add_paragraph()

    # ── Розділ 6: Q&A та зміни до ТД ─────────────────────────────────────────
    qa_analysis = analysis.get('qa_analysis') or {}
    _add_heading(doc, '6. Q&A ТА ЗМІНИ ДО ТД', level=2)

    if qa_analysis and not qa_analysis.get('analysis_skipped'):
        qa_table = doc.add_table(rows=4, cols=2)
        qa_table.style = 'Table Grid'

        qa_rows = [
            ('ТД містить зміни:', 'Так' if qa_analysis.get('td_has_amendments') else 'Ні'),
            ('Файл змін:', qa_analysis.get('amendments_file', '—')),
            ('Всього питань:', str(qa_analysis.get('questions_total', '—'))),
            ('Питань з відповіддю:', str(qa_analysis.get('questions_answered', '—'))),
        ]

        for i, (label, value) in enumerate(qa_rows):
            cells = qa_table.rows[i].cells
            cells[0].text = label
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].bold = True

        # Ключові уточнення
        clarifications = qa_analysis.get('key_clarifications') or []
        if clarifications:
            doc.add_paragraph()
            _add_heading(doc, 'Ключові уточнення від замовника:', level=3)
            for c in clarifications:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"[{c.get('topic', '—')}] ")
                run.bold = True
                p.add_run(c.get('customer_answer_summary', '—'))

        # Питання без відповіді
        unanswered = qa_analysis.get('unanswered_questions') or []
        if unanswered:
            doc.add_paragraph()
            _add_heading(doc, '⚠️ Питання без відповіді:', level=3)
            for q in unanswered:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(q.get('title', str(q)))

        # Статус оскарження після змін
        new_deadline = qa_analysis.get('new_appeal_deadline')
        if new_deadline:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f'Новий дедлайн оскарження змін: {new_deadline}')
            run.bold = True
            run.font.color.rgb = COLOR_RED
    else:
        doc.add_paragraph('Питань та відповідей не знайдено або аналіз пропущено.')

    doc.add_paragraph()

    # ── Розділ 7: Профіль замовника ───────────────────────────────────────────
    customer_profile = analysis.get('customer_profile') or {}
    _add_heading(doc, '7. ПРОФІЛЬ ЗАМОВНИКА', level=2)

    if customer_profile:
        oopz_count = customer_profile.get('oopz_decisions_count', 0)
        p = doc.add_paragraph()
        run = p.add_run(f"Рішень ООПЗ: {oopz_count}")
        if oopz_count and oopz_count > 3:
            run.font.color.rgb = COLOR_RED
        status = customer_profile.get('implementation_status', '')
        if status == 'partial':
            doc.add_paragraph(
                'Примітка: Повний профіль замовника в розробці (очікується інтеграція '
                'з рішеннями ООПЗ та Юконтрол).',
            )
    else:
        doc.add_paragraph('Профіль замовника відсутній.')

    doc.add_paragraph()

    # ── Розділ 8: Аналіз договору (якщо є) ───────────────────────────────────
    if contract_analysis:
        _add_heading(doc, '8. АНАЛІЗ ПРОЕКТУ ДОГОВОРУ', level=2)

        p = doc.add_paragraph()
        run = p.add_run(contract_analysis.get('appeal_deadline_note', ''))
        run.bold = True
        run.font.color.rgb = COLOR_RED

        doc.add_paragraph(f"Файл договору: {contract_analysis.get('contract_file', '—')}")

        balance = contract_analysis.get('overall_balance')
        risk = contract_analysis.get('risk_level')

        ct_table = doc.add_table(rows=3, cols=2)
        ct_table.style = 'Table Grid'

        ct_rows = [
            ('Баланс договору:', _balance_ua(balance)),
            ('Рівень ризику:', _risk_ua(risk)),
            ('Договір приєднання:', 'Так' if contract_analysis.get('adhesion_contract') else 'Ні'),
        ]
        for i, (label, value) in enumerate(ct_rows):
            cells = ct_table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
            cells[0].paragraphs[0].runs[0].bold = True
            if label == 'Баланс договору:':
                for run in cells[1].paragraphs[0].runs:
                    run.font.color.rgb = _risk_color(risk)

        doc.add_paragraph()
        summary = contract_analysis.get('summary', '')
        if summary:
            p = doc.add_paragraph()
            run = p.add_run('Висновок по договору: ')
            run.bold = True
            p.add_run(summary)

    doc.add_paragraph()

    # ── Розділ 9: Метадані аналізу ────────────────────────────────────────────
    metadata = analysis.get('metadata') or {}
    _add_heading(doc, '9. МЕТАДАНІ АНАЛІЗУ', level=2)

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'

    meta_rows = [
        ('Модель аналізу:', metadata.get('model_analysis', '—')),
        ('Модель класифікації:', metadata.get('model_classification', '—')),
        ('Тривалість аналізу:', f"{metadata.get('analysis_duration_sec', '—')} сек."),
        ('Промпти завантажені:', 'Так' if metadata.get('prompts_loaded') else 'Ні (заглушки)'),
    ]
    oopz_used = analysis.get('oopz_context_used') or []
    if oopz_used:
        meta_rows.append(('ООПЗ рішення використано:', ', '.join(oopz_used)))

    # Перебудувати таблицю з правильною кількістю рядків
    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.style = 'Table Grid'
    for i, (label, value) in enumerate(meta_rows):
        cells = meta_table.rows[i].cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].bold = True

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# 2. documents_checklist.docx — 6 колонок
# ═══════════════════════════════════════════════════════════════════════════════

def _build_documents_checklist(analysis: dict) -> Document:
    """
    Будує documents_checklist.docx — чеклист документів для тендерної пропозиції.

    6 колонок:
      № | Назва документа | Пункт/Розділ ТД | Цитата з ТД | Примітка ⚠️ | Нотатки користувача
    """
    doc = _setup_document()
    _add_doc_header(doc, analysis, 'ЧЕКЛИСТ ДОКУМЕНТІВ ДЛЯ ТЕНДЕРНОЇ ПРОПОЗИЦІЇ')

    _add_heading(doc, 'Перелік документів для подання у складі пропозиції', level=2)

    p = doc.add_paragraph(
        'Колонка "Нотатки" — залиште порожньою для ваших записів при підготовці пропозиції.'
    )
    p.runs[0].font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    document_blocks = analysis.get('document_blocks') or []

    if not document_blocks:
        doc.add_paragraph('Перелік документів не сформовано (аналіз у процесі розробки).')
        return doc

    global_counter = 1

    for block in document_blocks:
        if not isinstance(block, dict):
            continue
        # Заголовок блоку
        block_name = block.get('block_name', f'Блок {block.get("block_id", "?")}')
        td_ref = block.get('td_reference', '')
        heading_text = f'{block_name}'
        if td_ref:
            heading_text += f'  ({td_ref})'
        _add_heading(doc, heading_text, level=3)

        items = [i for i in (block.get('items') or []) if isinstance(i, dict)]
        if not items:
            doc.add_paragraph('Позиції відсутні.')
            continue

        # Таблиця 6 колонок
        table = doc.add_table(rows=1 + len(items), cols=6)
        table.style = 'Table Grid'

        # Заголовки
        col_headers = ['№', 'Назва документа', 'Пункт/Розділ ТД', 'Цитата з ТД', 'Примітка ⚠️', 'Нотатки']
        col_widths = [Cm(1), Cm(4), Cm(3), Cm(5), Cm(3), Cm(3)]

        for j, (h, w) in enumerate(zip(col_headers, col_widths)):
            cell = table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, '1F497D')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Встановити ширину колонки через XML
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w.pt * 20)))  # EMU → twips (1pt = 20twips)
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        for i, item in enumerate(items):
            row = table.rows[i + 1]

            num_cell = row.cells[0]
            num_cell.text = str(global_counter)
            global_counter += 1

            row.cells[1].text = str(item.get('document', item.get('doc_name', '—')))
            row.cells[2].text = str(item.get('td_reference', item.get('td_section', '—')))
            row.cells[3].text = str(item.get('quote', item.get('td_quote', '—')))

            # Примітка з ризик-маркером
            note = item.get('note', '')
            risk_flag = item.get('risk_flag', False)
            is_hidden = item.get('is_hidden', False)
            law_violation = item.get('law_violation', False)

            note_parts = []
            if is_hidden:
                note_parts.append('⚠️ ПРИХОВАНА вимога')
            if law_violation:
                law_ref = item.get('law_reference', '')
                note_parts.append(f'⚠️ Можливе порушення: {law_ref}' if law_ref else '⚠️ Можливе порушення')
            if risk_flag and not is_hidden and not law_violation:
                note_parts.append('⚠️ ' + note if note else '⚠️ Зверніть увагу')
            elif note and not note_parts:
                note_parts.append(note)

            note_cell = row.cells[4]
            note_cell.text = '\n'.join(note_parts) if note_parts else (note or '')
            if note_parts:
                for run in note_cell.paragraphs[0].runs:
                    run.font.color.rgb = COLOR_ORANGE

            # Нотатки — порожня
            row.cells[5].text = ''

            # Підфарбувати рядки з ризиком
            if risk_flag or is_hidden or law_violation:
                _set_cell_bg(row.cells[0], 'FFF2CC')
                _set_cell_bg(row.cells[1], 'FFF2CC')

        doc.add_paragraph()

    # Підсумок
    total_items = global_counter - 1
    _add_heading(doc, f'Всього позицій: {total_items}', level=2)

    hidden_count = sum(
        1 for block in document_blocks
        if isinstance(block, dict)
        for item in (block.get('items') or [])
        if isinstance(item, dict) and (item.get('is_hidden') or item.get('risk_flag'))
    )
    if hidden_count:
        p = doc.add_paragraph(f'⚠️ Позицій з ризик-маркером: {hidden_count}')
        p.runs[0].font.color.rgb = COLOR_ORANGE

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# 3. contract_analysis.docx
# ═══════════════════════════════════════════════════════════════════════════════

def _build_contract_analysis_doc(analysis: dict, contract_analysis: dict) -> Document:
    """
    Будує contract_analysis.docx — юридичний аналіз проекту договору.

    Структура:
      - КРИТИЧНИЙ ДЕДЛАЙН (жирно, червоно)
      - Загальна оцінка
      - Підстави для оскарження умов договору
      - Порушення ЦКУ / ЗУ 922 / КМУ №668
      - Односторонні умови
      - Приховані ризики
      - Висновок
    """
    doc = _setup_document()
    _add_doc_header(doc, analysis, 'ЮРИДИЧНИЙ АНАЛІЗ ПРОЕКТУ ДОГОВОРУ')

    # ── КРИТИЧНИЙ ДЕДЛАЙН (виділений!) ────────────────────────────────────────
    appeal_deadline_note = contract_analysis.get('appeal_deadline_note', '')
    if appeal_deadline_note:
        p = doc.add_paragraph()
        run = p.add_run(f'⚠️ КРИТИЧНО: {appeal_deadline_note}')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_RED

    p2 = doc.add_paragraph(
        'Зміни до умов договору можна ініціювати ВИКЛЮЧНО шляхом оскарження умов '
        'тендерної документації до ООПЗ в строк, зазначений вище. '
        'Після цієї дати — оскарження умов договору неможливе.'
    )
    p2.runs[0].font.color.rgb = COLOR_ORANGE
    doc.add_paragraph()

    # ── Розділ 1: Загальна оцінка ─────────────────────────────────────────────
    _add_heading(doc, '1. ЗАГАЛЬНА ОЦІНКА ДОГОВОРУ', level=2)

    balance = contract_analysis.get('overall_balance')
    risk_level = contract_analysis.get('risk_level')
    adhesion = contract_analysis.get('adhesion_contract', False)
    adhesion_note = contract_analysis.get('adhesion_note', '')

    ov_table = doc.add_table(rows=4, cols=2)
    ov_table.style = 'Table Grid'

    ov_rows = [
        ('Файл договору:', contract_analysis.get('contract_file', '—')),
        ('Договір приєднання (ст.634 ЦКУ):', 'Так' if adhesion else 'Ні'),
        ('Баланс умов:', _balance_ua(balance)),
        ('Рівень ризику для переможця:', _risk_ua(risk_level)),
    ]

    for i, (label, value) in enumerate(ov_rows):
        cells = ov_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(cells[0], 'D9E1F2')

        if label == 'Баланс умов:':
            for run in cells[1].paragraphs[0].runs:
                run.font.color.rgb = _risk_color(risk_level)
        elif label == 'Рівень ризику для переможця:':
            for run in cells[1].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = _risk_color(risk_level)

    if adhesion_note:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('Коментар: ')
        run.bold = True
        p.add_run(adhesion_note)

    doc.add_paragraph()

    # ── Розділ 2: Підстави для оскарження умов договору ──────────────────────
    appeal_grounds = contract_analysis.get('appeal_grounds') or []
    _add_heading(doc, f'2. ПІДСТАВИ ДЛЯ ОСКАРЖЕННЯ УМОВ ДОГОВОРУ ({len(appeal_grounds)} шт.)', level=2)

    if appeal_grounds:
        ag_table = doc.add_table(rows=1 + len(appeal_grounds), cols=5)
        ag_table.style = 'Table Grid'

        ag_headers = ['№', 'Порушення', 'Правова підстава', 'Пункт договору', 'Запропонована вимога']
        for j, h in enumerate(ag_headers):
            cell = ag_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, 'FCE4D6')

        for i, g in enumerate(appeal_grounds):
            row = ag_table.rows[i + 1]
            row.cells[0].text = str(i + 1)
            row.cells[1].text = str(g.get('ground', '—'))
            row.cells[2].text = str(g.get('legal_basis', g.get('law_basis', '—')))
            row.cells[3].text = str(g.get('contract_clause', '—'))
            row.cells[4].text = str(g.get('suggested_claim', '—'))
    else:
        doc.add_paragraph('Підстав для оскарження умов договору не виявлено.')

    doc.add_paragraph()

    # ── Розділ 3: Порушення законодавства ─────────────────────────────────────
    law_violations = contract_analysis.get('law_violations') or []
    _add_heading(doc, f'3. ПОРУШЕННЯ ЗАКОНОДАВСТВА ({len(law_violations)} шт.)', level=2)

    if law_violations:
        lv_table = doc.add_table(rows=1 + len(law_violations), cols=5)
        lv_table.style = 'Table Grid'

        lv_headers = ['Стаття', 'Опис порушення', 'Пункт договору', 'Ступінь', 'Вплив']
        for j, h in enumerate(lv_headers):
            cell = lv_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, 'FCE4D6')

        for i, v in enumerate(law_violations):
            row = lv_table.rows[i + 1]
            row.cells[0].text = str(v.get('article', '—'))
            row.cells[1].text = str(v.get('description', '—'))
            row.cells[2].text = str(v.get('contract_clause', '—'))
            severity = v.get('severity', '—')
            row.cells[3].text = severity
            if severity in ('high', 'critical'):
                for run in row.cells[3].paragraphs[0].runs:
                    run.font.color.rgb = COLOR_RED
            row.cells[4].text = str(v.get('impact', ''))
    else:
        doc.add_paragraph('Прямих порушень законодавства не виявлено.')

    doc.add_paragraph()

    # ── Розділ 4: Односторонні умови ─────────────────────────────────────────
    one_sided = contract_analysis.get('one_sided_conditions') or []
    _add_heading(doc, f'4. ОДНОСТОРОННІ УМОВИ НА ШКОДУ ПЕРЕМОЖЦЮ ({len(one_sided)} шт.)', level=2)

    if one_sided:
        os_table = doc.add_table(rows=1 + len(one_sided), cols=3)
        os_table.style = 'Table Grid'

        os_headers = ['Умова', 'Пункт договору', 'Вплив на переможця']
        for j, h in enumerate(os_headers):
            cell = os_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, 'FFF2CC')

        for i, c in enumerate(one_sided):
            row = os_table.rows[i + 1]
            row.cells[0].text = str(c.get('condition', '—'))
            row.cells[1].text = str(c.get('contract_clause', '—'))
            row.cells[2].text = str(c.get('impact', '—'))
    else:
        doc.add_paragraph('Односторонніх умов не виявлено.')

    doc.add_paragraph()

    # ── Розділ 5: Приховані ризики ─────────────────────────────────────────────
    hidden_risks = contract_analysis.get('hidden_risks') or []
    _add_heading(doc, f'5. ПРИХОВАНІ ФІНАНСОВІ ТА ОПЕРАЦІЙНІ РИЗИКИ ({len(hidden_risks)} шт.)', level=2)

    if hidden_risks:
        hr_table = doc.add_table(rows=1 + len(hidden_risks), cols=3)
        hr_table.style = 'Table Grid'

        hr_headers = ['Ризик', 'Пункт договору', 'Фінансовий вплив']
        for j, h in enumerate(hr_headers):
            cell = hr_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, 'E2EFDA')

        for i, r in enumerate(hidden_risks):
            row = hr_table.rows[i + 1]
            row.cells[0].text = str(r.get('risk', '—'))
            row.cells[1].text = str(r.get('contract_clause', '—'))
            row.cells[2].text = str(r.get('financial_impact', '—'))
    else:
        doc.add_paragraph('Прихованих ризиків не виявлено.')

    doc.add_paragraph()

    # ── Розділ 6: Загальний висновок ──────────────────────────────────────────
    _add_heading(doc, '6. ВИСНОВОК', level=2)

    summary = contract_analysis.get('summary', '')
    if summary:
        doc.add_paragraph(summary)

    p = doc.add_paragraph()
    run = p.add_run(f'РІВЕНЬ РИЗИКУ ДЛЯ ПЕРЕМОЖЦЯ: {_risk_ua(risk_level).upper()}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _risk_color(risk_level)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# 4. winner_checklist.docx
# ═══════════════════════════════════════════════════════════════════════════════

def _build_winner_checklist(analysis: dict, contract_analysis: Optional[dict]) -> Document:
    """
    Будує winner_checklist.docx — чеклист дій переможця.

    Структура:
      Розділ 1 — Документи протягом 4 днів після повідомлення про намір
      Розділ 2 — Документи при підписанні договору
      Розділ 3 — Зобов'язання під час виконання (важливо знати до підписання)
    """
    doc = _setup_document()
    _add_doc_header(doc, analysis, 'ЧЕКЛИСТ ПЕРЕМОЖЦЯ')

    warning_p = doc.add_paragraph(
        '⚠️ Цей документ призначений ВИКЛЮЧНО для переможця торгів. '
        'Не включається до складу тендерної пропозиції.'
    )
    warning_p.runs[0].bold = True
    warning_p.runs[0].font.color.rgb = COLOR_ORANGE
    doc.add_paragraph()

    winner_checklist = analysis.get('winner_checklist') or {}
    category = analysis.get('category', '')

    # Читаємо winner items з document_blocks (primary source — LLM-генерований контент)
    _winner_4days: list = []
    _winner_signing: list = []
    for _block in (analysis.get('document_blocks') or []):
        if not isinstance(_block, dict):
            continue
        _stage = _block.get('stage', '')
        for _item in (_block.get('items') or []):
            if not isinstance(_item, dict):
                continue
            _doc_name = _item.get('document') or _item.get('doc_name') or '—'
            _note = _item.get('note') or _item.get('td_quote') or ''
            if _note == 'None':
                _note = ''
            if _stage == 'winner_4days':
                _winner_4days.append({
                    'action': _doc_name,
                    'source': '—',
                    'deadline_days': 4,
                    'note': _note,
                })
            elif _stage == 'winner_signing':
                _winner_signing.append({
                    'action': _doc_name,
                    'source': '—',
                    'deadline': 'До підписання',
                    'note': _note,
                })

    # ── Розділ 1: Документи протягом 4 днів ───────────────────────────────────
    _add_heading(doc, 'РОЗДІЛ 1. Документи протягом 4 днів після повідомлення про намір', level=2)

    # Пріоритет: document_blocks → winner_checklist JSON → hardcoded default
    step1_from_json = winner_checklist.get('step1_4days') or []
    if _winner_4days:
        all_step1 = _winner_4days
    elif step1_from_json:
        all_step1 = step1_from_json
    else:
        all_step1 = [{
            'action': 'Витяг з ІАС МВС (відсутність судимості, корупційних правопорушень, '
                      'дитячої праці) на керівника учасника',
            'source': 'Сайт МВС / особисто',
            'deadline_days': 4,
            'note': 'Обов\'язково — ст.17 ЗУ 922-VIII',
        }]

    s1_headers = ['№', 'Дія / Документ', 'Де отримати', 'Строк', 'Примітка']
    s1_table = doc.add_table(rows=1 + len(all_step1), cols=5)
    s1_table.style = 'Table Grid'

    for j, h in enumerate(s1_headers):
        cell = s1_table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, '1F497D')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, item in enumerate(all_step1):
        row = s1_table.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = str(item.get('action', '—'))
        row.cells[2].text = str(item.get('source', '—'))
        deadline = item.get('deadline_days')
        row.cells[3].text = f"{deadline} дні(в)" if deadline else str(item.get('deadline', '—'))
        row.cells[4].text = str(item.get('note', ''))

    doc.add_paragraph()

    # ── Розділ 2: Документи при підписанні договору ───────────────────────────
    _add_heading(doc, 'РОЗДІЛ 2. Документи при підписанні договору', level=2)

    # Пріоритет: document_blocks → winner_checklist JSON → hardcoded default
    step2_from_json = winner_checklist.get('step2_contract_signing') or []
    if _winner_signing:
        all_step2 = _winner_signing
    elif step2_from_json:
        all_step2 = step2_from_json
    else:
        # Fallback — тільки універсальне забезпечення + АВК-5 лише для будівництва
        _CONSTRUCTION_CATEGORIES = {'building_works', 'maintenance_works'}
        all_step2 = [{
            'action': 'Перерахувати забезпечення виконання договору',
            'source': 'Власний рахунок → IBAN замовника',
            'deadline': 'До підписання',
            'note': '',
        }]
        if category in _CONSTRUCTION_CATEGORIES:
            all_step2.append({
                'action': 'Кошторисна документація у машинозчитувальному форматі (.word та .imd)',
                'source': 'АВК-5 або аналог',
                'deadline': 'До підписання',
                'note': 'Формат АВК-5 обов\'язковий',
            })

    s2_table = doc.add_table(rows=1 + len(all_step2), cols=5)
    s2_table.style = 'Table Grid'

    for j, h in enumerate(s1_headers):  # Самі заголовки
        cell = s2_table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, '1F497D')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, item in enumerate(all_step2):
        row = s2_table.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = str(item.get('action', '—'))
        row.cells[2].text = str(item.get('source', '—'))
        row.cells[3].text = str(item.get('deadline', 'До підписання'))
        row.cells[4].text = str(item.get('note', ''))

    doc.add_paragraph()

    # ── Розділ 3: Зобов'язання під час виконання ──────────────────────────────
    _add_heading(doc, 'РОЗДІЛ 3. Зобов\'язання під час виконання договору (важливо знати ДО підписання)', level=2)

    step3_items = winner_checklist.get('step3_contract_obligations') or []

    # Додаємо приховані ризики з contract_analysis як зобов'язання (якщо є)
    contract_obligations = []
    if contract_analysis:
        for risk in (contract_analysis.get('hidden_risks') or []):
            risk_text = risk.get('risk', '')
            contract_clause = risk.get('contract_clause', '')
            if risk_text:
                contract_obligations.append({
                    'obligation': risk_text,
                    'frequency': 'разово/регулярно',
                    'note': f"Договір {contract_clause}" if contract_clause else '⚠️ Перевірте договір',
                })

    all_step3 = step3_items + contract_obligations

    if all_step3:
        s3_table = doc.add_table(rows=1 + len(all_step3), cols=4)
        s3_table.style = 'Table Grid'

        s3_headers = ['№', 'Зобов\'язання', 'Частота', 'Примітка']
        for j, h in enumerate(s3_headers):
            cell = s3_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, '1F497D')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i, item in enumerate(all_step3):
            row = s3_table.rows[i + 1]
            row.cells[0].text = str(i + 1)
            row.cells[1].text = str(item.get('obligation', '—'))
            frequency = item.get('frequency', '—')
            freq_map = {
                'weekly': 'Щотижня',
                'monthly': 'Щомісяця',
                'once': 'Одноразово',
            }
            row.cells[2].text = freq_map.get(frequency, frequency)
            row.cells[3].text = str(item.get('note', ''))
    else:
        doc.add_paragraph(
            'Специфічних операційних зобов\'язань не виявлено. '
            'Перевірте договір самостійно.'
        )

    doc.add_paragraph()

    # ── Нагадування ───────────────────────────────────────────────────────────
    _add_heading(doc, 'ВАЖЛИВІ НАГАДУВАННЯ', level=2)

    reminders = [
        'Дедлайн 4 днів відраховується від дати повідомлення про намір укласти договір у Prozorro.',
        'Після підписання договору — внести забезпечення виконання (якщо вимагається) до першого платежу.',
        'Зберігайте копії всіх поданих документів.',
    ]

    if contract_analysis:
        deadline_note = contract_analysis.get('appeal_deadline_note', '')
        if deadline_note:
            reminders.insert(0, f'⚠️ {deadline_note}')

    for reminder in reminders:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(reminder)
        if reminder.startswith('⚠️'):
            run.bold = True
            run.font.color.rgb = COLOR_RED

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# 5. chronological_checklist.docx — хронологічний чеклист (Variant B сортування)
# ═══════════════════════════════════════════════════════════════════════════════

def _natural_sort_key(s: str) -> list:
    """Натуральне сортування: числа порівнюються як числа, текст — як рядки."""
    return [int(p) if p.isdigit() else p.lower() for p in re.split(r'(\d+)', s or '')]


def _td_ref_group(td_ref: str) -> int:
    """
    Variant B групування для хронологічного сортування:
      0 = невідомий/порожній td_reference (першим)
      1 = основна частина ТД (Розділ/Розд.) — не Додаток
      2 = Додатки (Додаток/Дод.) та Договір
    """
    ref_lower = (td_ref or '').strip().lower()
    if not ref_lower:
        return 0
    if any(m in ref_lower for m in ('додаток', 'дод.', 'дод ', 'договір', 'договор')):
        return 2
    if any(m in ref_lower for m in ('розділ', 'розд.', 'розд ')):
        return 1
    return 0


def _build_chronological_checklist(analysis: dict) -> Document:
    """
    Будує chronological_checklist.docx — всі документи зведені у порядку ТД.

    Variant B сортування:
      Група 0: невідомий розділ (першим)
      Група 1: основна частина (Розділи) — природне сортування
      Група 2: Додатки та Договір — природне сортування

    7 колонок:
      № | Назва документа | Блок | Пункт/Розділ ТД | Цитата з ТД | Прихована | Ризик
    """
    doc = _setup_document()
    _add_doc_header(doc, analysis, 'ЧЕКЛИСТ ДОКУМЕНТІВ — ХРОНОЛОГІЧНИЙ ПОРЯДОК ТД')

    _add_heading(doc, 'Всі вимоги у порядку розділів та додатків тендерної документації', level=2)

    p = doc.add_paragraph(
        'Сортування: спочатку основна частина ТД (Розділи), потім Додатки та Договір. '
        'Хронологічний порядок дозволяє звіряти чеклист з текстом ТД послідовно.'
    )
    p.runs[0].font.color.rgb = COLOR_GRAY
    doc.add_paragraph()

    document_blocks = analysis.get('document_blocks') or []
    if not document_blocks:
        doc.add_paragraph('Перелік документів не сформовано.')
        return doc

    # Збираємо всі items разом з назвою блоку
    flat_items = []
    for block in document_blocks:
        if not isinstance(block, dict):
            continue
        block_name = block.get('block_name', f'Блок {block.get("block_id", "?")}')
        for item in (block.get('items') or []):
            if not isinstance(item, dict):
                continue
            flat_items.append({**item, '_block_name': block_name})

    if not flat_items:
        doc.add_paragraph('Позиції відсутні.')
        return doc

    # Variant B сортування
    flat_items.sort(key=lambda it: (
        _td_ref_group(it.get('td_reference', '')),
        _natural_sort_key(it.get('td_reference', '')),
    ))

    # Таблиця 7 колонок
    table = doc.add_table(rows=1 + len(flat_items), cols=7)
    table.style = 'Table Grid'

    col_headers = ['№', 'Назва документа', 'Блок', 'Пункт/Розділ ТД', 'Цитата з ТД', 'Прихована', 'Ризик']
    col_widths = [Cm(0.8), Cm(4), Cm(3), Cm(3), Cm(5), Cm(1.5), Cm(1.5)]

    for j, (h, w) in enumerate(zip(col_headers, col_widths)):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, '1F497D')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(w.pt * 20)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    for i, item in enumerate(flat_items):
        row = table.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = str(item.get('document', item.get('doc_name', '—')))
        row.cells[2].text = str(item.get('_block_name', '—'))
        row.cells[3].text = str(item.get('td_reference', item.get('td_section', '—')))
        row.cells[4].text = str(item.get('quote', item.get('td_quote', '—')))

        is_hidden = item.get('is_hidden', False)
        risk_flag = item.get('risk_flag', False)

        hidden_cell = row.cells[5]
        hidden_cell.text = '⚠️ Так' if is_hidden else 'Ні'
        if is_hidden:
            for run in hidden_cell.paragraphs[0].runs:
                run.font.color.rgb = COLOR_ORANGE
                run.bold = True

        risk_cell = row.cells[6]
        risk_cell.text = '⚠️ Так' if risk_flag else 'Ні'
        if risk_flag:
            for run in risk_cell.paragraphs[0].runs:
                run.font.color.rgb = COLOR_RED
                run.bold = True

        if is_hidden or risk_flag:
            _set_cell_bg(row.cells[0], 'FFF2CC')
            _set_cell_bg(row.cells[1], 'FFF2CC')

    # Підсумок
    doc.add_paragraph()
    total = len(flat_items)
    hidden_count = sum(1 for it in flat_items if it.get('is_hidden'))
    risk_count = sum(1 for it in flat_items if it.get('risk_flag'))

    _add_heading(doc, f'Всього позицій: {total}', level=2)
    if hidden_count:
        p = doc.add_paragraph(f'⚠️ Прихованих вимог: {hidden_count}')
        p.runs[0].font.color.rgb = COLOR_ORANGE
    if risk_count:
        p = doc.add_paragraph(f'⚠️ З ризик-маркером: {risk_count}')
        p.runs[0].font.color.rgb = COLOR_RED

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# Збереження analysis.json
# ═══════════════════════════════════════════════════════════════════════════════

def _save_analysis_json(
    output_dir: Path,
    analysis: dict,
    contract_analysis: Optional[dict],
) -> Path:
    """
    Зберігає повний analysis.json у папці виводу.

    Структура JSON відповідає схемі з CLAUDE.md (блок ВИХІДНІ ДАНІ).
    Додає contract_analysis та analyzed_at до кореневого рівня.
    """
    analysis_out = dict(analysis)

    # Додаємо/оновлюємо поля верхнього рівня
    analysis_out['analyzed_at'] = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')

    if contract_analysis is not None:
        analysis_out['contract_analysis'] = contract_analysis
    elif 'contract_analysis' not in analysis_out:
        analysis_out['contract_analysis'] = None

    json_path = output_dir / 'analysis.json'
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(analysis_out, f, ensure_ascii=False, indent=2, default=str)

    logger.info("report_formatter: analysis.json збережено → %s", json_path)
    return json_path


# ═══════════════════════════════════════════════════════════════════════════════
# Головна функція
# ═══════════════════════════════════════════════════════════════════════════════

def format_reports(
    analysis: dict,
    contract_analysis: Optional[dict],
    output_dir: Path,
) -> dict:
    """
    Форматує результати аналізу у 5 DOCX файли та зберігає analysis.json.

    Args:
        analysis:          Результат base_analyzer.analyze() — повний словник аналізу ТД.
        contract_analysis: Результат contract_analyzer.analyze() — аналіз договору або None.
        output_dir:        Шлях до папки виводу (наприклад закупівлі/UA-XXXX/analysis/).
                           Папка буде створена якщо не існує.

    Returns:
        Словник з шляхами до збережених файлів:
          {
            'analysis_json': Path,
            'analysis_report': Path,
            'documents_checklist': Path,
            'contract_analysis_doc': Path | None,
            'winner_checklist': Path,
            'chronological_checklist': Path,
            'success': bool,
            'errors': list[str],
          }
    """
    start_time = time.time()
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    result = {
        'analysis_json': None,
        'analysis_report': None,
        'documents_checklist': None,
        'contract_analysis_doc': None,
        'winner_checklist': None,
        'chronological_checklist': None,
        'success': False,
        'errors': [],
    }

    public_id = analysis.get('public_id', 'unknown')
    logger.info(
        "report_formatter: старт форматування. public_id=%s, output_dir=%s",
        public_id, output_dir,
    )

    # ── 1. analysis.json ──────────────────────────────────────────────────────
    try:
        json_path = _save_analysis_json(output_dir, analysis, contract_analysis)
        result['analysis_json'] = json_path
    except Exception as exc:
        error_msg = f"analysis.json: {exc}"
        logger.error("report_formatter: помилка збереження %s", error_msg)
        result['errors'].append(error_msg)

    # ── 2. analysis_report.docx ───────────────────────────────────────────────
    try:
        doc = _build_analysis_report(analysis, contract_analysis)
        report_path = output_dir / 'analysis_report.docx'
        doc.save(str(report_path))
        result['analysis_report'] = report_path
        logger.info("report_formatter: analysis_report.docx збережено → %s", report_path)
    except Exception as exc:
        error_msg = f"analysis_report.docx: {exc}"
        logger.error("report_formatter: помилка %s", error_msg, exc_info=True)
        result['errors'].append(error_msg)

    # ── 3. documents_checklist.docx ───────────────────────────────────────────
    try:
        doc = _build_documents_checklist(analysis)
        checklist_path = output_dir / 'documents_checklist.docx'
        doc.save(str(checklist_path))
        result['documents_checklist'] = checklist_path
        logger.info("report_formatter: documents_checklist.docx збережено → %s", checklist_path)
    except Exception as exc:
        error_msg = f"documents_checklist.docx: {exc}"
        logger.error("report_formatter: помилка %s", error_msg, exc_info=True)
        result['errors'].append(error_msg)

    # ── 4. contract_analysis.docx (тільки якщо є contract_analysis) ──────────
    if contract_analysis:
        try:
            doc = _build_contract_analysis_doc(analysis, contract_analysis)
            contract_doc_path = output_dir / 'contract_analysis.docx'
            doc.save(str(contract_doc_path))
            result['contract_analysis_doc'] = contract_doc_path
            logger.info(
                "report_formatter: contract_analysis.docx збережено → %s", contract_doc_path
            )
        except Exception as exc:
            error_msg = f"contract_analysis.docx: {exc}"
            logger.error("report_formatter: помилка %s", error_msg, exc_info=True)
            result['errors'].append(error_msg)
    else:
        logger.info("report_formatter: contract_analysis відсутній — contract_analysis.docx пропущено")

    # ── 5. winner_checklist.docx ──────────────────────────────────────────────
    try:
        doc = _build_winner_checklist(analysis, contract_analysis)
        winner_path = output_dir / 'winner_checklist.docx'
        doc.save(str(winner_path))
        result['winner_checklist'] = winner_path
        logger.info("report_formatter: winner_checklist.docx збережено → %s", winner_path)
    except Exception as exc:
        error_msg = f"winner_checklist.docx: {exc}"
        logger.error("report_formatter: помилка %s", error_msg, exc_info=True)
        result['errors'].append(error_msg)

    # ── 6. chronological_checklist.docx ───────────────────────────────────────
    try:
        doc = _build_chronological_checklist(analysis)
        chrono_path = output_dir / 'chronological_checklist.docx'
        doc.save(str(chrono_path))
        result['chronological_checklist'] = chrono_path
        logger.info(
            "report_formatter: chronological_checklist.docx збережено → %s", chrono_path
        )
    except Exception as exc:
        error_msg = f"chronological_checklist.docx: {exc}"
        logger.error("report_formatter: помилка %s", error_msg, exc_info=True)
        result['errors'].append(error_msg)

    # ── Підсумок ──────────────────────────────────────────────────────────────
    duration = round(time.time() - start_time, 1)
    result['success'] = len(result['errors']) == 0

    logger.info(
        "report_formatter: завершено за %.1f сек. public_id=%s, success=%s, errors=%d",
        duration, public_id, result['success'], len(result['errors']),
    )

    return result
