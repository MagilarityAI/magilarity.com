"""
file_extractor.py — витяг тексту з файлів тендерної документації.

Підтримувані формати: .docx, .pdf, .xlsx, .xls, .txt
БЕЗ OCR. Зображення пропускаються. PDF без копійованого тексту → None.
"""
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Ключові слова для пошуку файлу проекту договору (case-insensitive)
_CONTRACT_KEYWORDS = [
    'договір підряду',
    'договор підряду',
    'договір про закупівлю',
    'додаток 3',
    'додаток №3',
    'contract',
    'dogovir',
    'проект договору',
]


# ---------------------------------------------------------------------------
# Внутрішні функції витягу за форматом
# ---------------------------------------------------------------------------

def _extract_docx(file_path: Path) -> Optional[str]:
    """Витяг тексту з .docx через python-docx."""
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts = []

        # Параграфи
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Таблиці
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = ' | '.join(cells)
                if row_text.strip():
                    parts.append(row_text)

        if not parts:
            return None
        return '\n'.join(parts)

    except Exception as e:
        logger.error(f"Помилка витягу docx '{file_path.name}': {e}")
        return None


def _extract_pdf(file_path: Path) -> Optional[str]:
    """Витяг тексту з copyable PDF. pdfplumber → fallback pypdf."""
    # Спроба через pdfplumber
    try:
        import pdfplumber

        with pdfplumber.open(str(file_path)) as pdf:
            parts = [
                page.extract_text()
                for page in pdf.pages
                if page.extract_text()
            ]

        if parts:
            return '\n'.join(parts)
        return None

    except ImportError:
        logger.warning(
            "pdfplumber не встановлено, використовую pypdf як fallback"
        )
    except Exception as e:
        logger.warning(
            f"pdfplumber не зміг обробити '{file_path.name}': {e} — "
            "перемикаюсь на pypdf"
        )

    # Fallback — pypdf
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        parts = [
            page.extract_text()
            for page in reader.pages
            if page.extract_text()
        ]

        if parts:
            return '\n'.join(parts)
        return None

    except Exception as e:
        logger.error(f"Помилка витягу pdf '{file_path.name}': {e}")
        return None


def _extract_xlsx(file_path: Path) -> Optional[str]:
    """Витяг тексту з .xlsx / .xls через openpyxl."""
    wb = None
    try:
        import openpyxl

        wb = openpyxl.load_workbook(
            str(file_path), read_only=True, data_only=True
        )
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Аркуш: {sheet_name} ===")

            for row in ws.iter_rows(values_only=True):
                cells = [
                    str(cell).strip()
                    for cell in row
                    if cell is not None and str(cell).strip()
                ]
                if cells:
                    parts.append(' | '.join(cells))

        if not parts:
            return None
        return '\n'.join(parts)

    except Exception as e:
        logger.error(f"Помилка витягу xlsx '{file_path.name}': {e}")
        return None

    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:
                pass


def _extract_txt(file_path: Path) -> Optional[str]:
    """Витяг тексту з .txt."""
    try:
        text = file_path.read_text(encoding='utf-8', errors='replace').strip()
        return text if text else None
    except Exception as e:
        logger.error(f"Помилка читання txt '{file_path.name}': {e}")
        return None


# ---------------------------------------------------------------------------
# Публічні функції
# ---------------------------------------------------------------------------

def extract_text(file_path: Path) -> Optional[str]:
    """
    Витягує текст з файлу тендерної документації.

    Визначає формат за суфіксом і делегує відповідній внутрішній функції.
    Формати без підтримки → повертає None (без виключень).

    Args:
        file_path: Шлях до файлу.

    Returns:
        Рядок з текстом або None якщо текст не вдалось витягти.
    """
    suffix = file_path.suffix.lower()

    if suffix == '.docx':
        return _extract_docx(file_path)
    elif suffix == '.pdf':
        return _extract_pdf(file_path)
    elif suffix in ('.xlsx', '.xls'):
        return _extract_xlsx(file_path)
    elif suffix == '.txt':
        return _extract_txt(file_path)
    else:
        logger.debug(
            f"Непідтримуваний формат '{suffix}' для файлу '{file_path.name}'"
        )
        return None


def extract_all_documents(
    docs_dir: Path,
    exclude_amendments: bool = False,
) -> dict:
    """
    Витягує текст з усіх файлів у директорії docs_dir.

    Args:
        docs_dir:          Директорія з файлами тендерної документації.
        exclude_amendments: Якщо True — файли-зміни ([ЗМІНИ]*) пропускаються
                            (не включаються в результат).

    Returns:
        Словник виду:
        {
            "filename.docx": {
                "path": "/абс/шлях/filename.docx",
                "text": "витягнутий текст" | None,
                "size": 12345,          # байти
                "is_amendment": False,
                "extraction_failed": False,
            },
            ...
        }
    """
    result: dict = {}

    if not docs_dir.exists():
        logger.error(f"Директорія не існує: {docs_dir}")
        return result

    for file_path in sorted(docs_dir.iterdir()):
        if not file_path.is_file():
            continue

        is_amendment = file_path.name.startswith('[ЗМІНИ]')

        if exclude_amendments and is_amendment:
            logger.debug(f"Пропускаємо файл-зміну: {file_path.name}")
            continue

        logger.info(f"Обробка файлу: {file_path.name}")
        text = extract_text(file_path)

        result[file_path.name] = {
            'path': str(file_path),
            'text': text,
            'size': file_path.stat().st_size,
            'is_amendment': is_amendment,
            'extraction_failed': text is None,
        }

    return result


def find_contract_file(docs_dir: Path) -> Optional[Path]:
    """
    Шукає файл проекту договору в директорії за ключовими словами в назві.

    Пошук case-insensitive. Повертає перший знайдений файл.

    Args:
        docs_dir: Директорія з файлами тендерної документації.

    Returns:
        Path до файлу договору або None якщо не знайдено.
    """
    if not docs_dir.exists():
        logger.error(f"Директорія не існує: {docs_dir}")
        return None

    for file_path in sorted(docs_dir.iterdir()):
        if not file_path.is_file():
            continue

        name_lower = file_path.name.lower()
        for keyword in _CONTRACT_KEYWORDS:
            if keyword.lower() in name_lower:
                logger.info(
                    f"Знайдено файл договору: '{file_path.name}' "
                    f"(ключове слово: '{keyword}')"
                )
                return file_path

    logger.warning(
        f"Файл договору не знайдено в директорії: {docs_dir}"
    )
    return None
